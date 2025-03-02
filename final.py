import streamlit as st
import pandas as pd
import numpy as np
import joblib
import time
from streamlit_lottie import st_lottie
import requests
import base64
from pathlib import Path
import google.generativeai as genai
import matplotlib.pyplot as plt
import os
import warnings
import subprocess
import sys
import plotly.express as px # type: ignore
import folium # type: ignore
from streamlit_folium import folium_static # type: ignore
from io import BytesIO
from reportlab.lib.pagesizes import letter # type: ignore
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle # type: ignore
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle # type: ignore
from reportlab.lib import colors # type: ignore
from docx import Document # type: ignore
import markdown
from streamlit_mic_recorder import mic_recorder, speech_to_text

# Set page configuration
st.set_page_config(
    page_title="Carbon Footprint Calculator",
    page_icon="🌍",
    layout="wide",
    initial_sidebar_state="expanded"
)

background_image = """
<style>
[data-testid="stAppViewContainer"] {
    background-image: url("https://img.freepik.com/free-vector/spring-floral-watercolor-background-vector-green-with-leaf-illustration_53876-126350.jpg");
    background-size: cover;
    background-position: center;
    background-repeat: no-repeat;
}
</style>
"""
st.markdown(background_image, unsafe_allow_html=True)

# Function to load Lottie animations
def load_lottieurl(url):
    r = requests.get(url)
    if r.status_code != 200:
        return None
    return r.json()

# Function to calculate BMI and determine body type
def calculate_body_type(weight, height):
    # Calculate BMI
    bmi = weight / (height/100)**2
    
    # Determine body type based on BMI
    if bmi < 18.5:
        return "underweight"
    elif bmi < 25:
        return "normal"
    elif bmi < 30:
        return "overweight"
    else:
        return "obese"

#IMPACT DISPLAY FOR REPORT
def calculate_and_display_impact(co2_footprint_kg):
    """Calculates and displays the tree and sea level impact of a CO2 footprint."""

    # Tree Calculations
    peepal_sequestration = 30  # kg CO2 per tree
    mahogany_sequestration = 22.5  # kg CO2 per tree
    teak_sequestration = 40  # kg CO2 per tree
    annual_co2_kg = co2_footprint_kg * 12
    peepal_trees = round(co2_footprint_kg / peepal_sequestration)
    mahogany_trees = round(co2_footprint_kg / mahogany_sequestration)
    teak_trees = round(co2_footprint_kg / teak_sequestration)

    # 🌳 Tree Impact Message
    tree_message = f"""
    <div style="font-size:20px; text-align:left; padding:10px; line-height:1.6;">
    Just as your individual CO2 emissions contribute to rising atmospheric CO2, your carbon footprint also reflects the loss of vital tree cover.
    While your impact may seem small, multiplied across billions of individuals, it significantly reduces Earth's ability to absorb CO2.  
    Let’s visualize how many Indian trees we will lose anually at this rate:
    <br><br><br>
    🌍 <b>Your Carbon Footprint:</b> {co2_footprint_kg:.2f} kg CO₂e led to:
    <ul>
        <li>🌱 The destruction of <span style="color:#28a745;">{peepal_trees} Peepal trees</span> 🌿 — their shade & oxygen lost forever.</li>
        <li>🌳 The felling of <span style="color:#8B4513;">{mahogany_trees} Indian Mahogany trees</span> 🪵 — a silent and devastating loss.</li>
        <li>🚜 The clear-cutting of <span style="color:#A0522D;">{teak_trees} Indian Teak trees</span> 🌲 — leaving a permanent scar on nature.</li>
        <li> <a href="https://catchfoundation.in/blogs/best-native-trees-carbon-sequestration-india">🔖 Read More Here</a></li>
    </ul>
    </div>
    """

    # 🌊 Sea Level Impact Calculation
    sea_level_impact_mm = (co2_footprint_kg / 1000) * 0.00000008 * 1000  # Convert to micrometers
    decade_impact=sea_level_impact_mm*10

    # 🌊 Sea Level Impact Message
    sea_message = f"""
    <div style="font-size:20px; text-align:left; padding:10px; line-height:1.6;">
    Your CO₂e footprint, though seemingly small, has a real-world consequence. 
    <br><br><br>
    🌊 <b>Your Carbon Footprint:</b> {co2_footprint_kg:.2f} kg CO₂e led to:
    <ul>
        <li>🌎 An estimated sea level rise of <span style="color:#1E90FF;">{sea_level_impact_mm:.6f} micrometers</span> 🌊</li>
        <li>Over the next decade, your emissions contribute to an estimated {decade_impact:.4f} ppm rise in global CO2. </li>  
        <li> <a href="https://www.pnas.org/doi/10.1073/pnas.1216073110">🔖 Read More Here</a></li>


    </ul>
    Remember, a 100 ppm increase is linked to a 20-30 cm sea level rise, and your emissions accelerate us towards that, threatening coastal communities.
    </div>
    """

    st.markdown("""
    <div style="border: 2px solid #4682B4; border-radius: 10px; padding: 20px; margin: 20px 0; background-color: #f0f8ff;">
        <h3 style="color: #1E90FF; text-align: center;  padding-bottom: 10px; margin-bottom: 20px;">
            🌊 🌳 Environmental Impact
        </h3>
    """, unsafe_allow_html=True)

    st.markdown(f"<h4>🌳 Tree Impact:</h4>", unsafe_allow_html=True)
    tree_col1, tree_col2 = st.columns([1, 1])

    with tree_col1:
        st.markdown(tree_message, unsafe_allow_html=True)

    with tree_col2:
        st.image("giphy.gif", use_container_width=True)  
    
    st.markdown("<br><br>", unsafe_allow_html=True)

    st.markdown(f"<h4>🌊 Sea Level Impact:</h4>", unsafe_allow_html=True)
    st.markdown("<br><br>", unsafe_allow_html=True)

    sea_col1, sea_col2 = st.columns([1, 1])

    with sea_col1:
        st.image("sea.gif", use_container_width=True)  

    with sea_col2:
        st.markdown(sea_message, unsafe_allow_html=True)



    # 📌 Note on Estimates
    st.markdown("""
    <p><i>Note: These calculations are simplified estimates from well-researched resources to help visualize your impact. 
    Actual environmental effects depend on many complex factors.</i></p>
    """, unsafe_allow_html=True)

    st.markdown("</div>", unsafe_allow_html=True) 
    
      
def compare_to_global_average(monthly_co2_kg):
    """Compares a user's monthly CO2 footprint to the global average and displays it with styled output."""

    global_average_annual_kg = 4000  # 4 metric tons converted to kilograms
    annual_co2_kg = monthly_co2_kg * 12
    ratio = monthly_co2_kg / global_average_annual_kg

    if ratio > 1:
        message = f"""
        <div style="font-size:20px; text-align:left; padding:10px; line-height:1.6;">

        🔥 Your emissions are <span style="color:#ff5733;"><b>{ratio:.2f} times</b> higher</span> than the global average of 4000 kg CO₂e per year.  
        <br>
        <b>🌡️ This means a higher impact on climate change, increased deforestation, and greater responsibility to reduce your footprint.</b>
        </div>
        """
    elif ratio < 1:
        lower_ratio = 1 / ratio
        message = f"""
        <div style="font-size:20px; text-align:left; padding:10px; line-height:1.6;">
        ✅ Your emissions are <span style="color:#28a745;"><b>{lower_ratio:.2f} times lower</b></span> than the global average of 4000 kg CO₂e per year.  
        <br>
        <b>🌱 This is great! Keep maintaining a sustainable lifestyle and encourage others to do the same.</b>
        </div>
        """
    else:
        message = f"""
        <div style="font-size:20px; text-align:left; padding:10px; line-height:1.6;">
        ⚖️ Your emissions match the global average of 4000 kg CO₂e per year.
        <br>
        <b>🔄 While you're at an average level, consider ways to reduce further for a greener future!</b>
        </div>
        """
    st.markdown("<br><br>", unsafe_allow_html=True)

    st.markdown(f"<h4>👁️👁️ Individual Impact:</h4>", unsafe_allow_html=True)

    person_col1, person_col2 = st.columns([1, 1])
    with person_col1:

        st.image("person.gif", use_container_width=True)  

    with person_col2:

        st.markdown(message, unsafe_allow_html=True)
        df = pd.DataFrame({
            'Category': ['Your CO₂', 'Global Average'],
            'Emissions (kg)': [monthly_co2_kg, global_average_annual_kg]
        })

        df['Color'] = ['#ff5733' if monthly_co2_kg > global_average_annual_kg else '#28a745', '#4682B4']

        fig = px.bar(
            df, 
            x='Category', 
            y='Emissions (kg)', 
            text='Emissions (kg)', 
            color='Color',
            color_discrete_map="identity",  
        )

        fig.update_traces(
            texttemplate='%{text:.0f} kg', 
            textposition='inside'
        )
        fig.update_layout(
            title="Carbon Footprint Comparison",
            yaxis_title="Annual CO₂ Emissions (kg)",
            xaxis_title="",
            showlegend=False
        )

        st.plotly_chart(fig, use_container_width=True)

#PERSONALIZED SUGGESTIONS
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
genai.configure(api_key=GOOGLE_API_KEY)
model1 = genai.GenerativeModel('gemini-2.0-flash')

def generate_eco_suggestion(issue):
    """Generate a personalized eco-friendly suggestion using Google Gemini."""
    prompt = f"""
    Generate a suggestion for someone who {issue}.
    Format it as a problem statement followed by a specific actionable solution.
    Make it detailed and conversational but not more than 3 bullets under each 'what you can do'.
    Ensure no extra text is displayed before and after it.
    Problem should be summarised in 2 sentences.
    Each bullet should have 2-3 sentences for explanation.
    What you can do should have a pointer followed by ':' followed by explanation.
    Example format: "Problem: [brief issue]. What you can do: [specific action]"
    """
    
    response = model1.generate_content(prompt)
    return response.text.strip()
# ✅ Conversion factor: 1 hectare absorbs 180,000 kg CO₂
CO2_ABSORPTION_PER_HECTARE_KG = 180000  

FORESTS = {
    "Birik Forest, West Bengal, India": (26.979467, 88.428268),
    "Berambadi State Forest, Karnataka, India": (11.768967, 76.479446)
}

def co2_to_forest_area(co2_kg):
    """Convert CO₂ footprint (kg) to lost forest area (m²)."""
    return (co2_kg / CO2_ABSORPTION_PER_HECTARE_KG) * 10000  # Convert to m²

def generate_polygon(center_lat, center_lon, area_m2):
    """Generate a square polygon to visualize lost forest area."""
    side_length = (area_m2 ** 0.5) / 111320  # Convert meters to degrees

    lat_min = center_lat - side_length / 2
    lat_max = center_lat + side_length / 2
    lon_min = center_lon - side_length / 2
    lon_max = center_lon + side_length / 2

    return [(lat_min, lon_min), (lat_min, lon_max), (lat_max, lon_max), (lat_max, lon_min), (lat_min, lon_min)]

def visualize_forest_loss(prediction, forest_name):
    """Generate and display the map based on CO₂ footprint and chosen forest."""
    if forest_name not in FORESTS:
        st.error("Invalid forest selection!")
        return

    latitude, longitude = FORESTS[forest_name]  
    lost_area_m2 = co2_to_forest_area(prediction)  # Convert CO₂ to area lost
    polygon_coords = generate_polygon(latitude, longitude, lost_area_m2)

    zoom_level = 30 

    # Folium map
    m = folium.Map(location=[latitude, longitude], zoom_start=zoom_level)

    folium.Polygon(
        locations=polygon_coords,
        color="red",
        fill=True,
        fill_color="red",
        fill_opacity=0.5,
        popup=f"Forest Loss: {lost_area_m2:.2f} m²",
    ).add_to(m)

    folium_static(m)
    
# Function to load model and related components
@st.cache_resource
def load_model():
    try:
        model = joblib.load("carbon_model.pkl")
        encoders = joblib.load("encoders.pkl")
        scaler = joblib.load("scaler.pkl")
        return model, encoders, scaler
    except FileNotFoundError:
        st.error("Model files not found. Please make sure the trained model and preprocessing files exist.")
        return None, None, None
sidebar_background = """
<style>
[data-testid="stSidebar"]::before {
    content: "";
    position: absolute;
    top: 0;
    left: 0;
    width: 100%;
    height: 100%;
    background: linear-gradient(rgba(255, 255, 255, 0.5), rgba(255, 255, 255, 0.5)), 
                url("https://png.pngtree.com/background/20211215/original/pngtree-leaves-autumn-watercolor-golden-leaf-outline-background-pattern-picture-image_1461252.jpg");
    background-size: cover;
    background-position: center;
    background-repeat: no-repeat;
    z-index: -1;
}
</style>
"""
st.markdown(sidebar_background, unsafe_allow_html=True)


st.sidebar.image("eco.png", use_container_width=True)
st.sidebar.title("🌿 Carbon Footprint App")
page = st.sidebar.radio("Navigation", ["About", "Calculate Footprint", "Enhance Your Awareness", "Query and Resolve"])


# About page
if page == "About":
    st.markdown("""
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Poppins:wght@300;400;600;700&family=Montserrat:wght@400;500;700&display=swap');
        
        .title-gradient {
            font-family: 'Montserrat', sans-serif;
            font-weight: 700;
            font-size: 36px;
            background: linear-gradient(90deg, #2E7D32, #1976D2);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            padding-bottom: 10px;
            margin-bottom: 20px;
        }
        
        .header-style {
            font-family: 'Montserrat', sans-serif;
            font-size: 28px;
            font-weight: 700;
            color: #00695C;
            margin-top: 15px;
            margin-bottom: 15px;
        }
        
        .subheader-style {
            font-family: 'Montserrat', sans-serif;
            font-size: 22px;
            font-weight: 600;
            color: #1E88E5;
            margin-top: 20px;
            margin-bottom: 10px;
        }
        
        p, li {
            font-family: 'Poppins', sans-serif;
            font-size: 16px;
            line-height: 1.6;
            color: #37474F;
        }
        
        .highlight-text {
            background: rgba(255, 255, 255, 0.25);
            backdrop-filter: blur(10px);
            padding: 20px;
            border-radius: 12px;
            border-left: 5px solid #26A69A;
            box-shadow: 0 4px 8px rgba(0,0,0,0.1);
            margin: 20px 0;
        }
        
        .feature-box {
            background: rgba(255, 255, 255, 0.2);
            backdrop-filter: blur(8px);
            padding: 16px;
            border-radius: 10px;
            box-shadow: 0 3px 6px rgba(0,0,0,0.08);
            margin-bottom: 15px;
            border-top: 3px solid #FFC300;
            transition: transform 0.3s ease;
        }
        
        .feature-box:hover {
            transform: translateY(-5px);
        }
        
        .feature-title {
            font-family: 'Montserrat', sans-serif;
            font-weight: 600;
            color: #00796B;
            font-size: 18px;
            margin-bottom: 8px;
        }
        
        .eco-factor-box {
            background: linear-gradient(120deg, rgba(0, 77, 64, 0.9), rgba(0, 121, 107, 0.9));
            backdrop-filter: blur(12px);
            color: white;
            padding: 25px;
            border-radius: 12px;
            text-align: center;
            margin: 25px auto;
            box-shadow: 0 6px 12px rgba(0,0,0,0.15);
            max-width: 800px;
        }
        
        .eco-factor-box p {
            color: white;
            font-family: 'Poppins', sans-serif;
        }
        
        .eco-factor-title {
            font-family: 'Montserrat', sans-serif;
            font-weight: 700;
            font-size: 24px;
            color: white;
            margin-bottom: 15px;
            border-bottom: 2px solid #4DB6AC;
            padding-bottom: 10px;
        }
        
        .emoji-bullet {
            list-style-type: none;
            padding-left: 5px;
        }
        
        .emoji-bullet li {
            padding-left: 28px;
            position: relative;
            margin-bottom: 12px;
        }
        
        .emoji-bullet li:before {
            position: absolute;
            left: 0;
            top: 2px;
        }
        
        .info-banner {
            background: rgba(225, 245, 254, 0.7);
            backdrop-filter: blur(5px);
            border-left: 5px solid #03A9F4;
            padding: 15px;
            border-radius: 8px;
            font-family: 'Poppins', sans-serif;
            margin-top: 25px;
            display: flex;
            align-items: center;
        }
        
        .info-banner-text {
            margin-left: 15px;
            font-weight: 500;
            color: #01579B;
        }
    </style>
    """, unsafe_allow_html=True)

    st.markdown('<h1 class="title-gradient">Welcome to Eco X!</h1>', unsafe_allow_html=True)
    
    st.markdown("""
    <div class="highlight-text">            
        <p>Eco X is your personal climate companion, designed to make understanding and reducing your carbon footprint engaging and accessible. I believe that everyone has the power to make a difference, and Eco X provides the tools to do just that.</p>
    </div>
    """, unsafe_allow_html=True)
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.markdown("""
        <h3 class="subheader-style">📊 Why Choose Eco X?</h3>
        
        <div class="feature-box">
            <div class="feature-title">🚗 Track Your Daily Footprint</div>
            <ul class="emoji-bullet">
                <li>📱 Monitor your environmental impact through easy-to-use tracking tools</li>
                <li>🍽️ Input your diet choices, energy consumption patterns and more!</li>
                <li>🌵 Get your carbon footprint calculated and watch how it affects the world</li>
                <li>💡 Allow us to provide you personalized measures to improve your footprint journey</li>
            </ul>
        </div>
        
        <div class="feature-box">
            <div class="feature-title">📈 Become Aware of the Consequences</div>
            <ul class="emoji-bullet">
                <li>📰 Browse through the latest news scraped from credible sources</li>
                <li>🌿 Choose across multiple categories</li>
                <li>🔄 Also, get the latest links for articles and blogs</li>
                <li>🏆 You gain the knowledge to understand and potentially influence the future of our planet.</li>
            </ul>
        </div>
        
        <div class="feature-box">
            <div class="feature-title">🔍 Draft Reports</div>
            <ul class="emoji-bullet">
                <li>🔎 Query the system your concern</li>
                <li>🌱 Access the latest sustainability research</li>
                <li>📜 Convert them into well formatted reports</li>
                <li>🤝 Download and share for your prefered purposes.</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)
        
    with col2:
        st.image("co.gif", use_container_width=True)  #
        
        st.markdown("""
        <div class="highlight-text">
            <div class="feature-title">🌱 What is a Carbon Footprint?</div>
            <p>A carbon footprint represents the total greenhouse gases released by your lifestyle choices and activities.</p>
            <p><b>Why is it important?</b></p>
            <ul class="emoji-bullet">
                <li>🌡️ Contributes to global warming</li>
                <li>🌪️ Leads to extreme weather events</li>
                <li>🌊 Causes rising sea levels</li>
                <li>🌎 Affects ecosystems worldwide</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)
        
    
    st.markdown("""
    <div class="eco-factor-box">
        <h3 class="eco-factor-title">🌟 You are the X Factor!</h3>
        <p>At Eco X, I believe that individual actions, when combined, create significant positive change for our planet.</p>
        <p>Your choices and commitment to sustainability are the catalysts for a greener future. Every small decision matters!</p>
        <p>Join our community today and let's make a difference, one footprint at a time!</p>
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown("""
    <div class="info-banner">
        <span style="font-size: 24px;">💻</span>
        <span class="info-banner-text">Navigate to the 'Calculate' tab to determine your carbon footprint!</span>
    </div>
    """, unsafe_allow_html=True)
    
elif page == "Calculate Footprint":
    st.title("🍂 Calculate Your Carbon Footprint")
    st.write("Fill in the details below to estimate your carbon emissions.")
    
    model, encoders, scaler = load_model()
    
    if model is not None and encoders is not None and scaler is not None:
        # Initialize session state for storing form values
        if 'weight' not in st.session_state:
            st.session_state.weight = 70.0
            st.session_state.height = 170.0
            st.session_state.sex = "male"
            st.session_state.diet = "omnivore"
            st.session_state.shower = "daily"
        
        with st.form("carbon_footprint_form"):
            col_left, col_right = st.columns(2)
            
            with col_left:
                
                # Personal Information Section
                with st.container():
                    st.markdown("""
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Poppins:wght@300;400;600&family=Montserrat:wght@400;600&display=swap');

        .custom-box {
            border: 2px solid #FFC300;
            border-radius: 10px;
            padding: 20px;
            margin: 20px 0;
            background-color: rgba(255, 255, 255, 0.2);
            backdrop-filter: blur(5px); /* Subtle blur */
            font-family: 'Poppins', sans-serif;
        }

        .custom-box h3 {
            color: brown;
            text-align: center;
            padding-bottom: 10px;
            margin-bottom: 20px;
            font-family: 'Montserrat', sans-serif;
            font-weight: 600;
            font-size: 30px;
        }
    </style>

    <div class="custom-box">
        <h3>🕵️‍♀️ Personal Information</h3>
    </div>
""", unsafe_allow_html=True)
                    
                    weight = st.number_input("Weight (kg)", 30.0, 200.0, st.session_state.weight, step=0.1)
                    height = st.number_input("Height (cm)", 100.0, 250.0, st.session_state.height, step=0.1)
                    
                    # Calculate body type based on BMI
                    body_type = calculate_body_type(weight, height)
                    st.info(f"Calculated Body Type: {body_type.capitalize()}")
                    
                    sex = st.selectbox("Sex", ["male", "female"], index=0)
                    diet = st.selectbox("Diet", ["omnivore", "vegetarian", "vegan", "pescatarian"], index=0)
                    shower_frequency = st.selectbox("Shower Frequency", 
                                                  ["daily", "more frequently", "less frequently","twice a day"], index=0)
                    
                    st.markdown("</div>", unsafe_allow_html=True)
                
                # Transportation Section
                with st.container():
                    st.markdown("""
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Poppins:wght@300;400;600&family=Montserrat:wght@400;600&display=swap');

        .custom-box {
            border: 2px solid #FFC300;
            border-radius: 10px;
            padding: 20px;
            margin: 20px 0;
            background-color: rgba(255, 255, 255, 0.2);
            backdrop-filter: blur(5px); /* Subtle blur */
            font-family: 'Poppins', sans-serif;
        }

        .custom-box h3 {
            color: brown;
            text-align: center;
            padding-bottom: 10px;
            margin-bottom: 20px;
            font-family: 'Montserrat', sans-serif;
            font-weight: 600;
            font-size: 30px;
        }
    </style>

    <div class="custom-box">
        <h3>🚙 Transportation</h3>
    </div>
""", unsafe_allow_html=True)
                    
                    transport_mode = st.selectbox("Primary Transport Mode", 
                                                 ["public", "private", "walk/bicycle"])
                    vehicle_type = st.selectbox("Vehicle Type (if applicable)", 
                                               ["", "petrol", "diesel", "hybrid", "electric","lpg"])
                    vehicle_distance = st.number_input("Monthly Vehicle Distance (km)", 0, 10000, 500)
                    air_travel = st.selectbox("Frequency of Air Travel", 
                                             ["never", "rarely", "occasionally", "frequently"])
                    
                    st.markdown("</div>", unsafe_allow_html=True)
            
            with col_right:
                
                # Home Energy Section
                with st.container():
                    st.markdown("""
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Poppins:wght@300;400;600&family=Montserrat:wght@400;600&display=swap');

        .custom-box {
            border: 2px solid #FFC300;
            border-radius: 10px;
            padding: 20px;
            margin: 20px 0;
            background-color: rgba(255, 255, 255, 0.2);
            backdrop-filter: blur(5px); /* Subtle blur */
            font-family: 'Poppins', sans-serif;
        }

        .custom-box h3 {
            color: brown;
            text-align: center;
            padding-bottom: 10px;
            margin-bottom: 20px;
            font-family: 'Montserrat', sans-serif;
            font-weight: 600;
            font-size: 30px;
        }
    </style>

    <div class="custom-box">
        <h3>🏡 Home Energy</h3>
    </div>
""", unsafe_allow_html=True)
                    
                    heating_source = st.selectbox("Heating Energy Source", 
                                                ["electricity", "natural gas", "coal", "wood"])
                    energy_efficiency = st.selectbox("Energy Efficiency Measures", 
                                                   ["Yes", "No", "Sometimes"])
                    tv_pc_hours = st.number_input("TV/PC Use (hours/day)", 0, 24, 2)
                    internet_hours = st.number_input("Internet Use (hours/day)", 0, 24, 3)
                    
                    st.markdown("</div>", unsafe_allow_html=True)
                
                # Consumption & Waste Section
                with st.container():
                    st.markdown("""
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Poppins:wght@300;400;600&family=Montserrat:wght@400;600&display=swap');

        .custom-box {
            border: 2px solid #FFC300;
            border-radius: 10px;
            padding: 20px;
            margin: 20px 0;
            background-color: rgba(255, 255, 255, 0.2);
            backdrop-filter: blur(5px); /* Subtle blur */
            font-family: 'Poppins', sans-serif;
        }

        .custom-box h3 {
            color: brown;
            text-align: center;
            padding-bottom: 10px;
            margin-bottom: 20px;
            font-family: 'Montserrat', sans-serif;
            font-weight: 600;
            font-size: 30px;
        }
    </style>

    <div class="custom-box">
        <h3>🛒 Consumption and Waste</h3>
    </div>
""", unsafe_allow_html=True)
                    
                    grocery_bill = st.number_input("Monthly Grocery Bill ($)", 0, 1000, 200)
                    new_clothes = st.number_input("New Clothes Purchased Monthly", 0, 100, 5)
                    
                    waste_size = st.selectbox("Waste Bag Size", ["small", "medium", "large", "extra large"])
                    waste_count = st.number_input("Waste Bags Weekly", 0, 20, 2)
                    
                    recycling = st.multiselect("What do you recycle?", 
                                             ["Paper", "Plastic", "Metal", "Glass", "Electronics"])
                    cooking_with = st.multiselect("Cooking Appliances", 
                                                ["Stove", "Oven", "Microwave", "Air Fryer", "Slow Cooker"])
                    
                    st.markdown("</div>", unsafe_allow_html=True)
            
            # Social Activities Section 
            with st.container():
                st.markdown("""
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Poppins:wght@300;400;600&family=Montserrat:wght@400;600&display=swap');

        .custom-box {
            border: 2px solid #FFC300;
            border-radius: 10px;
            padding: 20px;
            margin: 20px 0;
            background-color: rgba(255, 255, 255, 0.2);
            backdrop-filter: blur(5px); /* Subtle blur */
            font-family: 'Poppins', sans-serif;
        }

        .custom-box h3 {
            color: brown;
            text-align: center;
            padding-bottom: 10px;
            margin-bottom: 20px;
            font-family: 'Montserrat', sans-serif;
            font-weight: 600;
            font-size: 30px;
        }
    </style>

    <div class="custom-box">
        <h3>🎭 Social Activities</h3>
    </div>
""", unsafe_allow_html=True)
                
                social_activity = st.selectbox("Social Activity Frequency", 
                                              ["never", "sometimes", "often"])
                
                st.markdown("</div>", unsafe_allow_html=True)
            
            st.markdown("""
            <style>
            div.stButton > button {
                background-color: #4CAF50;
                color: white;
                font-weight: bold;
                padding: 10px 24px;
                border-radius: 8px;
                border: none;
                width: 100%;
                margin-top: 20px;
            }
            div.stButton > button:hover {
                background-color: #45a049;
            }
            </style>
            """, unsafe_allow_html=True)
            
            submitted = st.form_submit_button("🔍 Calculate Carbon Footprint")
        
        # Process form submission
        if submitted:
            with st.spinner("Calculating your carbon footprint..."):
                # Prepare user input for prediction
                user_input = {
                    "Body Type": body_type,  # Use calculated body type
                    "Sex": sex,
                    "Diet": diet,
                    "How Often Shower": shower_frequency,
                    "Heating Energy Source": heating_source,
                    "Transport": transport_mode,
                    "Vehicle Type": vehicle_type if vehicle_type else "None",
                    "Social Activity": social_activity,
                    "Monthly Grocery Bill": grocery_bill,
                    "Frequency of Traveling by Air": air_travel,
                    "Vehicle Monthly Distance Km": vehicle_distance,
                    "Waste Bag Size": waste_size,
                    "Waste Bag Weekly Count": waste_count,
                    "How Long TV PC Daily Hour": tv_pc_hours,
                    "How Many New Clothes Monthly": new_clothes,
                    "How Long Internet Daily Hour": internet_hours,
                    "Energy efficiency": energy_efficiency,
                    "Recycling": str(recycling) if recycling else "[]",
                    "Cooking_With": str(cooking_with) if cooking_with else "[]"
                }
                
                # Create DataFrame
                input_df = pd.DataFrame([user_input])
                
                # Encode categorical features
                categorical_cols = [
                    "Body Type", "Sex", "Diet", "How Often Shower", "Heating Energy Source",
                    "Transport", "Vehicle Type", "Social Activity", "Frequency of Traveling by Air",
                    "Waste Bag Size", "Energy efficiency", "Recycling", "Cooking_With"
                ]
                
                numerical_cols = [
                    "Monthly Grocery Bill", "Vehicle Monthly Distance Km", "Waste Bag Weekly Count",
                    "How Long TV PC Daily Hour", "How Many New Clothes Monthly", "How Long Internet Daily Hour"
                ]
                
                # Apply encoders
                for col in categorical_cols:
                    if col in encoders:
                        # Handle unseen categories
                        try:
                            input_df[col] = encoders[col].transform(input_df[col])
                        except ValueError:
                            # If category not seen during training, use a default value
                            input_df[col] = 0
                
                # Apply scaler to numerical columns
                input_df[numerical_cols] = scaler.transform(input_df[numerical_cols])
                
                # Make prediction
                prediction = model.predict(input_df)[0]
                
                progress_bar = st.progress(0)
                for i in range(100):
                    time.sleep(0.01)
                    progress_bar.progress(i + 1)
                
                st.markdown(f"""
                <div style="background-color: #e6ffe6; border-radius: 10px; padding: 20px; margin: 20px 0; 
                      border: 2px solid #3cb371; text-align: center;">
                    <h2>Your Estimated Monthly Carbon Footprint: {prediction:.2f} kg CO₂e</h2>
                </div>
                """, unsafe_allow_html=True)
                
                # Evaluate the result
                if prediction < 1500:
                    st.success("🌱 Your carbon footprint is below average. Great job!")
                    lottie_success = load_lottieurl("https://assets9.lottiefiles.com/packages/lf20_touohxv0.json")
                    if lottie_success:
                        st_lottie(lottie_success, height=200)
                elif prediction < 2500:
                    st.warning("⚠️ Your carbon footprint is about average. There's room for improvement!")
                    lottie_average = load_lottieurl("https://assets1.lottiefiles.com/private_files/lf30_gcroxmjc.json")
                    if lottie_average:
                        st_lottie(lottie_average, height=200)
                else:
                    st.error("🔥 Your carbon footprint is above average. Consider making changes!")
                    lottie_high = load_lottieurl("https://assets6.lottiefiles.com/temp/lf20_dgjK9i.json")
                    if lottie_high:
                        st_lottie(lottie_high, height=200)
                
                # Visual gauge for carbon footprint
                gauge_html = f"""
                <div style="text-align: center; margin: 20px 0;">
                    <div style="width: 100%; background-color: #e0e0e0; border-radius: 10px; height: 20px;">
                        <div style="width: {min(100, prediction/30)}%; background-color: {'#4CAF50' if prediction < 1500 else '#FFA500' if prediction < 2500 else '#FF0000'}; height: 20px; border-radius: 10px;"></div>
                    </div>
                    <div style="display: flex; justify-content: space-between; font-size: 12px; margin-top: 5px;">
                        <span>Low Impact</span>
                        <span>Average</span>
                        <span>High Impact</span>
                    </div>
                </div>
                """
                st.markdown(gauge_html, unsafe_allow_html=True)
                
                # Call the function to calculate and display environmental impact
                calculate_and_display_impact(prediction)
                compare_to_global_average(prediction)
                
                col1, col2 = st.columns([2, 1])  

                with col1:
                 
                    resource_tabs = st.tabs(list(FORESTS.keys()))

                    for forest_name, tab in zip(FORESTS.keys(), resource_tabs):
                        with tab:
                            st.subheader(f"Impact on {forest_name}")
                            visualize_forest_loss(prediction, forest_name)

                with col2:
                    
                    lost_area_m2 = co2_to_forest_area(prediction)
                    st.markdown("<br><br><br>", unsafe_allow_html=True)
                    st.markdown("<br><br>", unsafe_allow_html=True)
                    st.markdown("<br><br><br>", unsafe_allow_html=True)

                    st.warning(f"🚨 **Your carbon footprint alone causes {lost_area_m2:.2f} m² of forest loss.**")
                    
                    st.markdown(
                        """
                        Even though this may look small on a map, imagine **billions of people** emitting CO₂ every day. 
                        The impact is **devastating**. 🌍🔥

                        **Every tree lost means less oxygen, more heat, and fewer homes for wildlife.**  
                        **What if there was a way to offset this loss?**  

                        👉 [🌱 **Click here to plant a tree and reduce your footprint** 🌱](https://onetreeplanted.org)
                        """,
                        unsafe_allow_html=True,
                    )

                
                
                # Suggestions
                st.markdown("""
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Poppins:wght@300;400;600&family=Montserrat:wght@400;600&display=swap');

        .custom-box {
            border: 2px solid #FFC300;
            border-radius: 10px;
            padding: 20px;
            margin: 20px 0;
            background-color: rgba(255, 255, 255, 0.2);
            backdrop-filter: blur(5px); /* Subtle blur */
            font-family: 'Poppins', sans-serif;
        }

        .custom-box h3 {
            color: brown;
            text-align: center;
            padding-bottom: 10px;
            margin-bottom: 20px;
            font-family: 'Montserrat', sans-serif;
            font-weight: 600;
            font-size: 30px;
        }
    </style>

    <div class="custom-box">
        <h3>🌟 Personalized Suggestions</h3>
    </div>
""", unsafe_allow_html=True)

                
                col1, col2 = st.columns(2)
                
                # Identify areas for improvement
                suggestion_triggers = []

                if transport_mode == "private" and vehicle_type in ["petrol", "diesel"]:
                    suggestion_triggers.append("drives a personal gasoline/diesel car regularly")
                    
                if tv_pc_hours + internet_hours > 10:
                    suggestion_triggers.append("spends more than 10 hours daily on screens")
                    
                if waste_count > 3:
                    suggestion_triggers.append("produces a large amount of non-recyclable waste")
                    
                if heating_source in ["coal", "wood"]:
                    suggestion_triggers.append("uses coal or wood for home heating")
                    
                if not recycling:
                    suggestion_triggers.append("doesn't recycle household waste")
                    
                if diet == "omnivore":
                    suggestion_triggers.append("consumes meat in their diet regularly")
                    
                if new_clothes > 10:
                    suggestion_triggers.append("purchases more than 10 new clothing items per year")

                # Generate dynamic suggestions
                with st.spinner("Generating personalized eco-friendly suggestions..."):
                    suggestions = []
                    for issue in suggestion_triggers:
                        try:
                            suggestion = generate_eco_suggestion(issue)
                            suggestions.append(suggestion)
                        except Exception as e:
                            st.error(f"Error generating suggestion: {str(e)}")
                            # Fallback to static suggestions if API fails
                            if "gasoline" in issue:
                                suggestions.append("Problem: Gas vehicles emit CO2. What you can do: Try carpooling or public transit twice a week.")
                            elif "screen" in issue:
                                suggestions.append("Problem: High screen time uses electricity. What you can do: Set device-free hours and use power-saving modes.")
                st.markdown(f"""
                        <style>
                        .feature-box {{
                            font-size: 18px; /* Adjust this size as needed */
                            font-weight: bold;
                            padding: 10px;
                            line-height: 1.6;
                        }}
                        </style>
                        """, unsafe_allow_html=True)
                with col1:
                    for i in range(0, len(suggestions), 2):
                        
                        st.markdown(f"""
                        <div class="feature-box">
                            <p>{suggestions[i]}</p>
                        </div>
                        """, unsafe_allow_html=True)

                with col2:
                    for i in range(1, len(suggestions), 2):
                        st.markdown(f"""
                        <div class="feature-box">
                            <p>{suggestions[i]}</p>
                        </div>
                        """, unsafe_allow_html=True)

                # If no suggestions were generated
                if not suggestions:
                    st.info("Great job! We don't have any specific suggestions for improvement based on your current habits.")

#ENHANCE AWARENESS
elif page == "Enhance Your Awareness":
    st.title("📚 Read to Succeed")
    
    resource_tabs = st.tabs(["Environmental News", "Articles & Blogs"])
    
    emoji_map = {
        "climate change": "🌡️",
        "carbon footprint": "👣",
        "sustainable living": "♻️",
        "green technology": "🔋",
        "renewable energy": "🌞",
        "eco-friendly living": "🌱"
    }

    category_labels = {
        "Climate Change": "climate change",
        "Carbon Footprint": "carbon footprint",
        "Sustainable Living": "sustainable living", 
        "Green Technology": "green technology",
        "Renewable Energy": "renewable energy",
        "Eco-Friendly Living": "eco-friendly living"
    }
    
    col1, col2 = st.columns([1, 3])
    with col1:
        selected_category_label = st.sidebar.selectbox("Select Category", list(category_labels.keys()))
        category = category_labels[selected_category_label]
        

    category_emoji = emoji_map.get(category, "🌍")
    
    # Environmental News Tab
    with resource_tabs[0]:
        st.subheader(f"Environmental News: {selected_category_label} {category_emoji}")
        
        st.markdown("""
        <style>
        .grid-item {
            border: 1px solid #e0e0e0;
            padding: 15px;
            margin-bottom: 15px;
            margin-top: 10px;
            border-radius: 5px;
            border-color: green;
            height: 325px;
            overflow: hidden;
        }

        .grid-item h3 {
            margin-bottom: 10px;
            font-size: 19px;
            word-wrap: break-word;
            white-space: normal;
        }

        .grid-item p {
            overflow: hidden;
            font-size: 16px;
            text-overflow: ellipsis;
            display: -webkit-box;
            -webkit-line-clamp: 3;
            -webkit-box-orient: vertical;
        }

        .grid-item a {
            color: #00bfff;
            text-decoration: none;
            font-weight: bold;
            transition: color 0.3s ease;
        }

        .grid-item a:hover {
            color: green;
            cursor: pointer;
        }
        
        .article-card {
            border: 1px solid #4caf50;
            border-radius: 8px;
            padding: 16px;
            margin-bottom: 16px;
            background-color: #f9f9f9;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            transition: transform 0.2s ease-in-out;
        }
        
        .article-card:hover {
            transform: translateY(-5px);
            box-shadow: 0 4px 8px rgba(0,0,0,0.15);
        }
        
        .article-title {
            color: #2e7d32;
            font-size: 18px;
            margin-bottom: 8px;
        }
        
        .article-snippet {
            color: #555;
            font-size: 14px;
            margin-bottom: 12px;
        }
        
        .article-link {
            color: #1976d2;
            text-decoration: none;
            font-weight: 500;
            font-size: 14px;
        }
        
        .article-link:hover {
            text-decoration: underline;
        }
        </style>
        """, unsafe_allow_html=True)

        def load_data():
            """
            Loads environmental news data from the CSV file.
            If the file doesn't exist, runs the scraping script to create it first.
            
            Returns:
                dict: A dictionary with category names as keys and dataframes as values.
            """
            categories = [cat for cat in category_labels.values()]
            
            output_csv = "cleaned_file.csv"
            
            # Check if the CSV exists
            if not os.path.exists(output_csv):
                st.warning("Data file not found. Running the scraper to collect articles. This might take a while...")
                
                try:
                    # Run the scraping script in the same directory
                    with st.spinner("Scraping news data... This may take several minutes..."):
                        result = subprocess.run([sys.executable, "new.py"], 
                                              capture_output=True, 
                                              text=True,
                                              check=True)
                        st.success("Scraping completed successfully!")
                        st.info(result.stdout)
                except subprocess.CalledProcessError as e:
                    st.error(f"Error running the scraper: {e}")
                    st.error(f"Error details: {e.stderr}")
                    return {}
                except Exception as e:
                    st.error(f"Unexpected error: {str(e)}")
                    return {}
                    
                if not os.path.exists(output_csv):
                    st.error(f"Error: File {output_csv} still not found after running the scraper.")
                    return {}
            
            try:
                df = pd.read_csv(output_csv)
                
                # Split data by category
                category_data = {}
                for cat in categories:
                    category_data[cat] = df[df['category'].str.lower() == cat.lower()]
                
                return category_data
            except Exception as e:
                st.error(f"Error loading data: {str(e)}")
                return {}
        
        with st.spinner("Loading Environmental News... 🌍"):
            category_data = load_data()
        
        # Check if data exists for the selected category
        if category in category_data and not category_data[category].empty:
            df_category = category_data[category]
            
            # Define the number of articles per page and total pages
            articles_per_page = 6
            max_articles = 30  
            df_category = df_category.head(max_articles)  # Select the latest 30 articles
            total_pages = min((len(df_category) + articles_per_page - 1) // articles_per_page, 5)
            
            if total_pages > 0:
                page_number = st.radio(
                    "Page Navigation",
                    range(1, total_pages + 1),
                    horizontal=True,
                    key="news_page_navigation",
                    label_visibility="collapsed"  
                )
                
                start_idx = (page_number - 1) * articles_per_page
                end_idx = start_idx + articles_per_page
                
                page_data = df_category.iloc[start_idx:end_idx]
                
                cols = st.columns(2)  
                for idx, row in page_data.iterrows():
                    col = cols[idx % 2]  
                    with col:
                        st.markdown(f'''
                        <div class="grid-item">
                            <h3>{row['title']}</h3>
                            <p>{row['subtitle']}</p>
                            <p><strong>Author</strong>: {row['author']}</p>
                            <a href="{row['link']}" target="_blank">Read Full Article</a>
                        </div>
                        ''', unsafe_allow_html=True)
            else:
                st.info(f"No articles to display for the {selected_category_label} category.")
        else:
            st.info(f"No articles found for the {selected_category_label} category. Please select another category or check if the data is loaded correctly.")
    
    # Articles & Blogs Tab
    with resource_tabs[1]:
        st.subheader(f"Articles & Blogs: {selected_category_label} {category_emoji}")
        
        # Function to search Google using the Custom Search API
        @st.cache_data(ttl=3600)  # Cache results for 1 hour
        def search_google(query, num_results=10):
            try:
                from googleapiclient.discovery import build
                
                API_KEY = os.getenv("GOOGLE_API_KEY")
                CSE_ID = os.getenv("GOOGLE_CSE_ID")
                
                if API_KEY == "API_KEY_NOT_FOUND" or CSE_ID == "CSE_ID_NOT_FOUND":
                    st.warning("Google Search API credentials not found. Please configure them in your Streamlit secrets.")
                    return []
                
                service = build("customsearch", "v1", developerKey=API_KEY)
                results = []
                
                res = service.cse().list(
                    q=query,
                    cx=CSE_ID,
                    num=num_results
                ).execute()
                
                for item in res.get("items", []):
                    results.append({
                        "title": item.get("title"),
                        "link": item.get("link"),
                        "snippet": item.get("snippet")
                    })
                    
                return results
                
            except Exception as e:
                st.error(f"Error searching Google: {str(e)}")
                return []
        
        with st.spinner(f"Searching for {selected_category_label} articles and blogs..."):
            search_query = f"environmental {category} articles blogs guides"
            search_results = search_google(search_query, num_results=10)
        
        if search_results:
            for i, result in enumerate(search_results):
                st.markdown(f"""
                <div class="article-card">
                    <div class="article-title">{result['title']}</div>
                    <div class="article-snippet">{result['snippet']}</div>
                    <a class="article-link" href="{result['link']}" target="_blank">Read more →</a>
                </div>
                """, unsafe_allow_html=True)
        else:
            # If no results or API not configured, show helpful resources and tips
            st.info("Couldn't fetch articles from search. Here are some general resources:")
            
            # Default
            resource_mapping = {}
            
            default_resources = [
                {"title": "United Nations Sustainable Development Goals", "link": "https://sdgs.un.org/goals", 
                 "desc": "Framework for peace and prosperity for people and the planet."},
                {"title": "World Wildlife Fund", "link": "https://www.worldwildlife.org/", 
                 "desc": "Leading organization in wildlife conservation and endangered species."}
            ]
            
            # Get resources for the current category or use defaults
            resources_to_show = resource_mapping.get(category, default_resources)
            
            for resource in resources_to_show:
                st.markdown(f"""
                <div class="article-card">
                    <div class="article-title">{resource['title']}</div>
                    <div class="article-snippet">{resource['desc']}</div>
                    <a class="article-link" href="{resource['link']}" target="_blank">Visit resource →</a>
                </div>
                """, unsafe_allow_html=True)
            
            st.markdown("""
            <div style="background-color: #f0f7ff; padding: 12px; border-radius: 5px; border-left: 4px solid #1976d2; margin-top: 20px;">
                <p style="margin: 0; font-size: 14px;">
                    <strong>Note:</strong> To enable live article search, add your Google Custom Search API key and Engine ID to your Streamlit secrets.toml file:
                </p>
                <pre style="background-color: #f5f5f5; padding: 8px; border-radius: 3px; margin-top: 8px; font-size: 12px;">
GOOGLE_API_KEY = "your_api_key_here"
GOOGLE_CSE_ID = "your_custom_search_engine_id_here"
                </pre>
            </div>
            """, unsafe_allow_html=True)


# QUERY AND REPORT
else:  


    st.markdown("""
    <style>
        .stButton button {
            width: 100%;
            border-radius: 5px;
            font-weight: bold;
        }
        .report-container {
            border: 1px solid #ddd;
            border-radius: 5px;
            padding: 20px;
            background-color: #f9f9f9;
            margin-top: 20px;
        }
        h1, h2, h3 {
            color: #1E3A8A;
        }
        .error-message {
            color: #CF142B;
            padding: 10px;
            background-color: #FFEEEE;
            border-radius: 5px;
            margin: 10px 0;
        }
        .carbon-section {
            background-color: #E6F7E6;
            border-left: 5px solid #28A745;
            padding: 10px;
            margin: 15px 0;
        }
    </style>
    """, unsafe_allow_html=True)

    # Initialize session state variables for report content
    if 'report_content' not in st.session_state:
        st.session_state.report_content = ""

    CLIMATE_KEYWORDS = [
        "climate", "carbon footprint", "emission", "greenhouse gas", "global warming", 
        "climate change", "renewable", "sustainability", "sustainable", "environment", 
        "environmental", "ecology", "ecological", "carbon", "green energy", "clean energy",
        "pollution", "co2", "methane", "fossil fuel", "temperature", "warming", "net zero",
        "climate crisis", "climate action", "carbon neutral", "carbon offset", "carbon credit"
    ]

    # Function to validate if query is climate-related
    def is_climate_related(query):
        query = query.lower()
        return any(keyword in query for keyword in CLIMATE_KEYWORDS)

    # Google search function
    def perform_google_search(query, num_results=5):
        try:
            api_key = os.environ.get("GOOGLE_SEARCH_API_KEY")
            cse_id = os.environ.get("GOOGLE_CSE_ID")
            
            if not api_key or not cse_id:
                return None
                
            url = f"https://www.googleapis.com/customsearch/v1"
            params = {
                "key": api_key,
                "cx": cse_id,
                "q": query,
                "num": num_results
            }
            
            response = requests.get(url, params=params)
            results = response.json()
            
            search_results = []
            if "items" in results:
                for item in results["items"]:
                    search_results.append({
                        "title": item.get("title", ""),
                        "link": item.get("link", ""),
                        "snippet": item.get("snippet", "")
                    })
            
            return search_results
        except Exception as e:
            st.error(f"Error in Google Search: {str(e)}")
            return None

    def generate_report(question, tone, min_words, max_words, additional_instructions="", specific_sections="", use_search=False, include_tables=True):
        try:
            search_results = None
            if use_search:
                with st.spinner("Searching for relevant information..."):
                    search_results = perform_google_search(question)
            
            sections_text = ""
            if specific_sections:
                sections_list = specific_sections.strip().split("\n")
                sections_text = "Include these specific sections:\n" + "\n".join([f"- {section}" for section in sections_list])
            
            prompt = f"""
            Generate a comprehensive {tone.lower()} report based on this climate or carbon footprint related question: '{question}'
            
            Requirements:
            - The report should have between {min_words} and {max_words} words
            - Include an executive summary/introduction
            - Include a conclusion section
            - MANDATORY: Include a dedicated section titled "Relation with Carbon Footprint" that specifically explains how the topic relates to carbon emissions, carbon footprint measurement, reduction strategies, and climate impact, even if the original query didn't explicitly mention this
            - Include a resources/references section with citations
            - Use proper headings and subheadings for organization
            - Use bullet points and numbered lists where appropriate
            {f"- Include descriptive tables where relevant data can be presented" if include_tables else ""}
            - DO NOT include image placeholders
            
            {sections_text}
            
            {additional_instructions}
            
            The report MUST focus on climate change, sustainability, or carbon footprint topics. Make sure the "Relation with Carbon Footprint" section is substantial (at least 150 words) and provides meaningful insights.
            
            Format the report in Markdown syntax. DO NOT begin or end your response with ```markdown tags or any other code block tags.
            """
            
            if search_results:
                prompt += "\n\nHere are some relevant search results you can use as references:\n"
                for i, result in enumerate(search_results, 1):
                    prompt += f"{i}. {result['title']} - {result['link']}\n   {result['snippet']}\n\n"
            
            # Initialize Gemini model
            model = genai.GenerativeModel('gemini-2.0-flash')
            
            with st.spinner("Generating your climate report..."):
                response = model.generate_content(prompt)
                report_content = response.text
                
                # Clean up markdown code block syntax if present
                if report_content.startswith("```markdown") or report_content.startswith("```"):
                    report_content = report_content.replace("```markdown", "", 1)
                    report_content = report_content.replace("```", "", 1)
                
                if report_content.endswith("```"):
                    report_content = report_content.rsplit("```", 1)[0]
                
                report_content = report_content.strip()
            
            return report_content
        
        except Exception as e:
            st.error(f"Error generating report: {str(e)}")
            return None

    def get_available_width(doc):
        return doc.width - doc.leftMargin - doc.rightMargin

    def create_pdf(content):
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
        styles = getSampleStyleSheet()
        
        # Update all styles to include proper wrapping
        for style_name in styles.byName:
            styles[style_name].wordWrap = 'CJK'
            styles[style_name].allowWidows = 0
            styles[style_name].allowOrphans = 0
        
        try:
            custom_heading1 = ParagraphStyle(
                name='CustomHeading1', 
                parent=styles['Heading1'],
                fontSize=16, 
                spaceAfter=12, 
                textColor=colors.blue,
                wordWrap='CJK'
            )
            custom_heading2 = ParagraphStyle(
                name='CustomHeading2', 
                parent=styles['Heading2'],
                fontSize=14, 
                spaceAfter=10, 
                textColor=colors.navy,
                wordWrap='CJK'
            )
            custom_heading3 = ParagraphStyle(
                name='CustomHeading3', 
                parent=styles['Heading3'],
                fontSize=12, 
                spaceAfter=8, 
                textColor=colors.darkblue,
                wordWrap='CJK'
            )
            carbon_section = ParagraphStyle(
                name='CarbonSection', 
                parent=styles['Normal'],
                fontSize=12, 
                spaceAfter=8, 
                backColor=colors.lightgreen, 
                borderColor=colors.green, 
                borderWidth=1, 
                borderPadding=5,
                wordWrap='CJK'
            )
        except KeyError:
            pass
        
        elements = []
        
        elements.append(Paragraph("Climate & Carbon Footprint Report", styles['Title']))
        elements.append(Spacer(1, 12))
        
        lines = content.split('\n')
        current_list = []
        in_list = False
        in_table = False
        table_data = []
        in_carbon_section = False
        available_width = get_available_width(doc)
        
        for line in lines:
            if "# Relation with Carbon Footprint" in line or "## Relation with Carbon Footprint" in line:
                in_carbon_section = True
                if in_list:
                    elements.append(Table([[bullet] for bullet in current_list]))
                    current_list = []
                    in_list = False
                try:
                    elements.append(Paragraph(line.replace('#', '').strip(), custom_heading1))
                except:
                    elements.append(Paragraph(line.replace('#', '').strip(), styles['Heading1']))
                elements.append(Spacer(1, 6))
                continue
            elif in_carbon_section and (line.startswith('# ') or line.startswith('## ')):
                in_carbon_section = False
            
            # Headings
            if line.startswith('# '):
                if in_list:
                    elements.append(Table([[bullet] for bullet in current_list]))
                    current_list = []
                    in_list = False
                try:
                    elements.append(Paragraph(line[2:], custom_heading1))
                except:
                    elements.append(Paragraph(line[2:], styles['Heading1']))
            elif line.startswith('## '):
                if in_list:
                    elements.append(Table([[bullet] for bullet in current_list]))
                    current_list = []
                    in_list = False
                try:
                    elements.append(Paragraph(line[3:], custom_heading2))
                except:
                    elements.append(Paragraph(line[3:], styles['Heading2']))
            elif line.startswith('### '):
                if in_list:
                    elements.append(Table([[bullet] for bullet in current_list]))
                    current_list = []
                    in_list = False
                try:
                    elements.append(Paragraph(line[4:], custom_heading3))
                except:
                    elements.append(Paragraph(line[4:], styles['Heading3']))
            
            # Lists
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                in_list = True
                current_list.append(line.strip()[2:])
            
            # Tables (simplified)
            elif line.strip().startswith('|') and not in_table:
                in_table = True
                table_data = [line.strip().split('|')[1:-1]]
            elif line.strip().startswith('|') and in_table:
                if not line.strip().startswith('|-'):  # Skip separator row
                    table_data.append(line.strip().split('|')[1:-1])
            elif in_table:
                # Table end
                in_table = False
                if table_data:
                    col_widths = [available_width/len(table_data[0])]*len(table_data[0])
                    table = Table(table_data, colWidths=col_widths)
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                        ('WORDWRAP', (0, 0), (-1, -1), True)
                    ]))
                    elements.append(table)
                    elements.append(Spacer(1, 12))
                
                if line.strip():
                    try:
                        style = carbon_section if in_carbon_section else styles['Normal']
                    except:
                        style = styles['Normal']
                    elements.append(Paragraph(line, style))
            
            # Regular paragraphs
            elif line.strip() and not in_list and not in_table:
                try:
                    style = carbon_section if in_carbon_section else styles['Normal']
                except:
                    style = styles['Normal']
                elements.append(Paragraph(line, style))
                elements.append(Spacer(1, 6))
        
        if in_list:
            elements.append(Table([[bullet] for bullet in current_list]))
        
        # Build the PDF
        doc.build(elements)
        buffer.seek(0)
        return buffer

    def create_docx(content):
        doc = Document()
        
        section = doc.sections[0]
        section.left_margin = 914400 // 8 
        section.right_margin = 914400 // 8
        
        doc.add_heading('Climate & Carbon Footprint Report', 0)
        
        lines = content.split('\n')
        current_list = []
        in_list = False
        in_table = False
        table_data = []
        in_carbon_section = False
        
        for line in lines:
            if "# Relation with Carbon Footprint" in line or "## Relation with Carbon Footprint" in line:
                in_carbon_section = True
                if in_list:
                    paragraph = doc.add_paragraph()
                    for item in current_list:
                        paragraph.add_run('• ' + item + '\n')
                    current_list = []
                    in_list = False
                heading = doc.add_heading(line.replace('#', '').strip(), 1)
                continue
            elif in_carbon_section and (line.startswith('# ') or line.startswith('## ')):
                in_carbon_section = False
            
            # Headings
            if line.startswith('# '):
                if in_list:
                    paragraph = doc.add_paragraph()
                    for item in current_list:
                        paragraph.add_run('• ' + item + '\n')
                    current_list = []
                    in_list = False
                doc.add_heading(line[2:], 1)
            elif line.startswith('## '):
                if in_list:
                    paragraph = doc.add_paragraph()
                    for item in current_list:
                        paragraph.add_run('• ' + item + '\n')
                    current_list = []
                    in_list = False
                doc.add_heading(line[3:], 2)
            elif line.startswith('### '):
                if in_list:
                    paragraph = doc.add_paragraph()
                    for item in current_list:
                        paragraph.add_run('• ' + item + '\n')
                    current_list = []
                    in_list = False
                doc.add_heading(line[4:], 3)
            
            # Lists
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                in_list = True
                current_list.append(line.strip()[2:])
            
            # Tables
            elif line.strip().startswith('|') and not in_table:
                in_table = True
                table_data = [line.strip().split('|')[1:-1]]
            elif line.strip().startswith('|') and in_table:
                if not line.strip().startswith('|-'):  
                    table_data.append(line.strip().split('|')[1:-1])
            elif in_table and table_data:
                # Table end
                in_table = False
                
                # Create the table
                if len(table_data) > 1:  
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    table.style = 'Table Grid'
                    
                    for i, row_data in enumerate(table_data):
                        for j, cell_data in enumerate(row_data):
                            if j < len(table.rows[i].cells):  
                                table.rows[i].cells[j].text = cell_data.strip()
                    
                    for cell in table.rows[0].cells:
                        cell.paragraphs[0].runs[0].bold = True
                    
                    table_width = section.page_width - section.left_margin - section.right_margin
                    for column in table.columns:
                        column.width = int(table_width / len(table_data[0]))
                    
                    doc.add_paragraph() 
                
                if line.strip():
                    para = doc.add_paragraph(line)
                    if in_carbon_section:
                        para.style = 'Quote'
            
            # Regular paragraphs
            elif line.strip() and not in_list and not in_table:
                para = doc.add_paragraph(line)
                if in_carbon_section:
                    para.style = 'Quote'
        
        if in_list:
            paragraph = doc.add_paragraph()
            for item in current_list:
                paragraph.add_run('• ' + item + '\n')
        
        # Save to a BytesIO object
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def create_download_link(buffer, filename, format_type):
        buffer.seek(0)
        b64 = base64.b64encode(buffer.read()).decode()
        
        mime_type = "application/pdf" if format_type == "PDF" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        href = f'<a href="data:{mime_type};base64,{b64}" download="{filename}" class="download-button">Download {format_type}</a>'
        
        return href

    # Main content area
    st.header("🌍 Generate Comprehensive Reports to Your Queries")
    st.markdown("Generate detailed climate and carbon footprint reports using Google's Gemini AI model")


    if "question" not in st.session_state:
        st.session_state.question = ""

    # 📝 **Text Input Field**
    question = st.text_area(
        "Enter your climate or carbon footprint-related question",
        value=st.session_state.question,  
        height=100,
        placeholder="Example: Tell me more about carbon footprint",
    )

    # 🎙️ **Voice-to-Text Below Text Input**
    text = speech_to_text(language="en", use_container_width=True, just_once=True, key="STT")

    if text:
        st.session_state.question = text
        st.rerun()  

    st.write(f"**Your Recording:** {st.session_state.question}")

    st.subheader("Report Configuration")

    col_style1, col_style2 = st.columns(2)

    with col_style1:
        tone = st.selectbox("Report Tone", 
                        ["Professional", "Academic", "Conversational", "Technical", "Simplified"])
        min_words = st.number_input("Minimum Words", min_value=300, max_value=2000, value=500, step=100)

    with col_style2:
        format_choice = st.selectbox("Download Format", ["PDF", "DOCX"])
        max_words = st.number_input("Maximum Words", min_value=500, max_value=3000, value=1000, step=100)

    include_tables = st.checkbox("Include data tables (if relevant)", value=True)
    use_search = st.checkbox("Use Google Search for enriched content", value=False)

    with st.expander("Advanced Options"):
        additional_instructions = st.text_area(
            "Additional Instructions (Optional)",
            placeholder="Example: Focus on economic impacts or Include recent technological advancements",
            height=100
        )
        
        specific_sections = st.text_area(
            "Additional Specific Sections to Include (Optional, one per line)",
            placeholder="Example:\nCurrent Market Analysis\nRegulatory Framework\nFuture Outlook",
            height=100
        )

    if st.button("Generate Report", type="primary"):
        if not st.session_state.question:
            st.error("Please enter a question or topic")
        elif not is_climate_related(st.session_state.question):
            st.error("This application only processes questions related to climate change and carbon footprint. Please modify your question.")
        else:
            report_content = generate_report(
                question=st.session_state.question,
                tone=tone,
                min_words=min_words,
                max_words=max_words,
                additional_instructions=additional_instructions,
                specific_sections=specific_sections,
                use_search=use_search,
                include_tables=include_tables
            )
            
            if report_content:
                st.session_state.report_content = report_content
                st.success("Climate report generated successfully!")

    if st.session_state.report_content:
        st.markdown("---")
        st.header("Generated Climate Report")
        
        preview_tab, download_tab = st.tabs(["Preview", "Download"])
        
        with preview_tab:
            st.markdown('<div class="report-container">', unsafe_allow_html=True)
            
            report_lines = st.session_state.report_content.split('\n')
            in_carbon_section = False
            formatted_report = []
            
            for line in report_lines:
                if "# Relation with Carbon Footprint" in line or "## Relation with Carbon Footprint" in line:
                    in_carbon_section = True
                    formatted_report.append('<div class="carbon-section">')
                    formatted_report.append(line)
                    continue
                elif in_carbon_section and (line.startswith('# ') or line.startswith('## ')):
                    formatted_report.append('</div>')
                    in_carbon_section = False
                    formatted_report.append(line)
                else:
                    formatted_report.append(line)
            
            if in_carbon_section:
                formatted_report.append('</div>')
            
            st.markdown('\n'.join(formatted_report))
            st.markdown('</div>', unsafe_allow_html=True)
        
        with download_tab:
            st.subheader("Download Options")
            
            if format_choice == "PDF":
                buffer = create_pdf(st.session_state.report_content)
                filename = "climate_report.pdf"
                st.markdown(
                    create_download_link(buffer, filename, "PDF"),
                    unsafe_allow_html=True
                )
                st.info("Click the link above to download your report as a PDF file.")
            else:  
                buffer = create_docx(st.session_state.report_content)
                filename = "climate_report.docx"
                st.markdown(
                    create_download_link(buffer, filename, "DOCX"),
                    unsafe_allow_html=True
                )
                st.info("Click the link above to download your report as a Word document.")